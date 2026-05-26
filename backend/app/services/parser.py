import re
import base64
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from markdown import markdown


class DocumentParser:
    """Parse different document types into structured text."""

    SUPPORTED_TYPES = {"pdf", "docx", "txt", "md", "markdown"}

    @classmethod
    def parse(cls, file_path: str, file_type: str) -> str:
        file_type = file_type.lower()
        if file_type == "pdf":
            return cls._parse_pdf(file_path)
        elif file_type == "docx":
            return cls._parse_docx(file_path)
        elif file_type in ("txt", "text"):
            return cls._parse_txt(file_path)
        elif file_type in ("md", "markdown"):
            return cls._parse_markdown(file_path)
        raise ValueError(f"Unsupported file type: {file_type}")

    @classmethod
    def _parse_pdf(cls, file_path: str) -> str:
        parts = []
        parts.append(cls._parse_pdf_text(file_path))
        parts.append(cls._parse_pdf_tables(file_path))
        return "\n\n".join(p for p in parts if p)

    @staticmethod
    def _parse_pdf_text(file_path: str) -> str:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        doc.close()
        return "\n".join(text_parts)

    @staticmethod
    def _parse_pdf_tables(file_path: str) -> str:
        table_texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        rows = []
                        for row in table:
                            cells = [str(c) if c else "" for c in row]
                            rows.append(" | ".join(cells))
                        if rows:
                            table_texts.append("\n".join(rows))
        return "\n\n".join(table_texts) if table_texts else ""

    @classmethod
    def extract_images_from_pdf(cls, file_path: str) -> list[bytes]:
        """Extract embedded images for chart recognition."""
        images = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                images.append(base_image["image"])
        doc.close()
        return images

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        doc = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()


def detect_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".text": "txt",
        ".md": "md",
        ".markdown": "md",
    }
    return mapping.get(ext, "unknown")
